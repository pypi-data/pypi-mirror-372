from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import tempfile
import json
import re
from groq import BadRequestError, AuthenticationError, RateLimitError, InternalServerError, APIError
from openai import OpenAI
from openai.types.chat import ChatCompletion
import tempfile

from pyparsing import line


def docx_to_json(docx_file_path):
    doc = Document(docx_file_path)

    if not doc.tables:
        text = []

        for para in doc.paragraphs:
            cleaned = para.text.strip()
            if cleaned:  # Skip empty lines
                text.append(cleaned)

        return "\n".join(text)

    table_data = []
    table = doc.tables[0]

    for row in table.rows:
        row_data = {}
        for i, cell in enumerate(row.cells):
            header = table.cell(0, i).text.strip()
            row_data[header] = cell.text.strip()
        table_data.append(row_data)

    return table_data


def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        txt = file.read()
    return txt


def import_file(file_dict):
    if file_dict['type'] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return docx_to_json(file_dict['datapath'])
    elif file_dict['type'] == "text/plain":
        return read_txt(file_dict['datapath'])
    else:
        return None


def count_pupils(transcript):
  # Regex für Sprecher
  sprecher_pattern = r'\b(LEHRER|S\d{2})\b(?=:)'

  # Alle Sprecher finden
  sprecher_liste = re.findall(sprecher_pattern, transcript)

  # Einzigartige Sprecher
  einzigartige_sprecher = set(sprecher_liste)

  # Prüfen, ob "LEHRER" vorkommt
  lehrer_vorhanden = "LEHRER" in einzigartige_sprecher

  # Sprecher ohne LEHRER
  sprecher_ohne_lehrer = {s for s in einzigartige_sprecher if s != "LEHRER"}

  # Logik für tatsächliche Anzahl ohne Lehrer
  if lehrer_vorhanden:
      anzahl_schueler = len(sprecher_ohne_lehrer)
  else:
      # Lehrer ist eine der S-Nummern → eine Person abziehen
      anzahl_schueler = max(len(einzigartige_sprecher) - 1, 0)  # keine negativen Zahlen

  return anzahl_schueler


def dialog_stats(transcript, lehrperson):
 # beitrag_pattern = r'(?:^|(?<=\s)|(?<=//))\b(' + lehrperson + r'|S\d{2})\b:\s*(.*)'
 # beitrag_pattern = r'\b(' + lehrperson + r'|S\d{2})\b:\s*(.*)'
 # beitrag_pattern = r'(?://\s*)?(?:\b(' + lehrperson + r'|S\d{2})\b):\s*(.*)'
 # beitrag_pattern = r'(?:^|//)?\b(' + lehrperson + r'|S\d{2})\b:\s*(.*)'
 # beitrag_pattern = rf'(?:^|//\s*|\s+){lehrperson}|S\d{2})\b:\s*(.*)'

# Extrahieren der Beiträge
  # 1. Einschübe aufsplitten
  text_split = re.sub(r"//(.*?)//", r"\n\1\n", transcript, flags=re.DOTALL)
  # 2. Regex für Beiträge
  beitrag_pattern = re.compile(rf"\b({lehrperson}|S\d{{2}})\b:\s*(.*)")
  beitraege = beitrag_pattern.findall(text_split)

  # DataFrame bauen
  df = pd.DataFrame(beitraege, columns=["Sprecher", "Beitrag"])

  # Wortanzahl je Beitrag berechnen
  df['Wortanzahl'] = df['Beitrag'].str.split().apply(len)

  # Zusammenfassung
  df_summary = df.groupby('Sprecher').agg(
      Anzahl_Beitraege=('Beitrag', 'count'),
      Gesamt_Woerter=('Wortanzahl', 'sum'),
      Durchschnitt_Woerter=('Wortanzahl', 'mean'),
      Median_Woerter=('Wortanzahl', 'median')
  ).reset_index()

    # 🟣 Schüler vs. Lehrer trennen
  df_lehrer = df_summary[df_summary['Sprecher'] == lehrperson]
  df_schueler = df_summary[df_summary['Sprecher'] != lehrperson]

  # Schüler zusammenfassen:
  schueler_summary = pd.DataFrame({
      'Sprecher': ['Schüler:innen'],
      'Anzahl_Beitraege': [df_schueler['Anzahl_Beitraege'].sum()],
      'Gesamt_Woerter': [df_schueler['Gesamt_Woerter'].sum()],
      'Durchschnitt_Woerter': [df_schueler['Gesamt_Woerter'].sum() / df_schueler['Anzahl_Beitraege'].sum()],
      'Median_Woerter': [df_schueler['Median_Woerter'].median()]
  })

  # Neu zusammenfügen:
  df_summary_neu = pd.concat([df_lehrer, schueler_summary], ignore_index=True)
  return df_summary_neu





def llm_analysis_groq(system_prompt, user_prompt, model, transcript, codebook, client):
    try:
        # Create chat completion object with JSON response format
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": user_prompt.format(transcript=transcript, codebook=codebook),
                }
            ],
            model=model,
            response_format={"type": "json_object"}
        )

        analysis_json_string = chat_completion.choices[0].message.content
        return analysis_json_string
    
    except BadRequestError as e:
        print(f"❗️BadRequestError (400): {str(e)}")
        return json.dumps({"error": "Bad request - Failed to generate JSON."})

    except AuthenticationError as e:
        print(f"❗️AuthenticationError (403): {str(e)}")
        return json.dumps({"error": "Authentication failed - Check API key or access rights."})

    except RateLimitError as e:
        print(f"❗️RateLimitError (429): {str(e)}")
        return json.dumps({"error": "Rate limit exceeded - Too many requests."})

    except InternalServerError as e:
        print(f"❗️InternalServerError (500+): {str(e)}")
        return json.dumps({"error": "Server error - Please try again later."})

    except APIError as e:
        print(f"❗️APIError: {str(e)}")
        return json.dumps({"error": f"API error: {str(e)}"})

    except Exception as e:
        print(f"❗️Unexpected error: {str(e)}")
        return json.dumps({"error": f"Unexpected error: {str(e)}"})
    

def llm_analysis_openai(
    system_prompt: str,
    user_prompt: str,
    model: str,
    transcript,
    codebook,
    client: OpenAI
) -> str:
    try:

        # Define schema for structured output
        schema = {
            "type": "object",
            "properties": {
                "analysis": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "#": {"type": "integer", "description": "Nummerierung"},
                            "Shortcode": {"type": "string", "description": "Der Shortcode"},
                            "Lehreräußerung (kurz)": {"type": "string", "description": "Die Äußerung"}
                        },
                        "required": ["#", "Shortcode", "Lehreräußerung (kurz)"],
                        "additionalProperties": False
                    },
                    "description": "Liste von Analyseobjekten"
                }
            },
            "required": ["analysis"],
            "additionalProperties": False
        }


        # Make the API call with structured output
        response = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt.format(transcript=transcript, codebook=codebook)}
            ],
            text={
                "format": {
                    "type": "json_schema",
                    "name": "analysis",
                    "schema": schema,
                    "strict": True
                }
            }
        )

        return response.output_text

    except Exception as e:
        print(f"❌ OpenAI API error: {e}")
        return json.dumps({"error": str(e)})



def count_teacher_impulses(df, teacher_name):
  anzahl = df.loc[df['Sprecher'] == teacher_name, 'Anzahl_Beitraege'].values[0]
  return anzahl

def remove_table_borders(table):
    tbl = table._tbl  # Access the XML element
    tblPr = tbl.tblPr

    tblBorders = tblPr.xpath('./w:tblBorders')
    if tblBorders:
        tblPr.remove(tblBorders[0])  # Remove existing borders

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'nil')  # 'nil' removes the line
        borders.append(edge_el)

    tblPr.append(borders)

def set_row_borders(row, top=False, bottom=False, left=False, right=False, size=12, color="000000", space="0"):
    for cell in row:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        def set_border(side):
            side_el = tcBorders.find(qn(f'w:{side}'))
            if side_el is None:
                side_el = OxmlElement(f'w:{side}')
                tcBorders.append(side_el)
            side_el.set(qn('w:val'), 'single')
            side_el.set(qn('w:sz'), str(size))       # border thickness
            side_el.set(qn('w:color'), color)        # hex color
            side_el.set(qn('w:space'), space)

        if top:
            set_border('top')
        if bottom:
            set_border('bottom')
        if left:
            set_border('left')
        if right:
            set_border('right')


def generate_report(
    output_path: str,
    group_name,
    pupil_count, 
    participation_quota, 
    quant_distribution_plot, 
    teacher_imp_count, 
    qual_distribution_plot
):
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as quant_distribution_plot_tmp:
        quant_distribution_plot.figure.savefig(quant_distribution_plot_tmp.name)
        quant_distribution_plot_tmp.close()      

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as qual_distribution_plot_tmp:
        qual_distribution_plot.figure.savefig(qual_distribution_plot_tmp.name)
        qual_distribution_plot_tmp.close()
    
    report = Document()

    report.sections[0].left_margin = Inches(0.5)
    report.sections[0].right_margin = Inches(0.5)
    report.sections[0].top_margin = Inches(0.5)
    report.sections[0].bottom_margin = Inches(0.5)

    report.add_heading('Ergebnisse Simulationstraining – Gruppe ' + group_name, 1)
    report.add_heading('1. Quantitative Verteilung der Gesprächsanteile', level=2)

    p = report.add_paragraph('Anzahl beteiligter Schüler:innen: ')
    p.add_run()
    p.add_run(' Beteiligungsquote: ')
    p.add_run(str(round(participation_quota, 1)))
    p.add_run('%)')

    report.add_paragraph('Tabelle')
    report.add_paragraph().add_run('Gesprächsbeiträge der Lehrperson und der Schüler:innen').italic=True

    table1 = report.add_table(rows=3, cols=3)
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Teilnehmende'
    hdr_cells[1].text = 'Lehrperson'
    hdr_cells[2].text = 'Schüler:innen'


    report.add_paragraph('Abbildung')
    report.add_paragraph().add_run('Gesprächsverteilung').italic=True

    report.add_picture(quant_distribution_plot_tmp.name, width=Inches(5))


    report.add_heading('2.	Qualitative Codierung der Impulse der Lehrperson, level 1', level=2)

    report.add_paragraph('Anzahl der Gesprächsimpulse: N = ' + str(teacher_imp_count))
    p1 = report.add_paragraph('Abbildung und Tabelle ')
    p1.add_run('Sprachliche Impulse der Lehrperson').italic = True

    report.add_picture(qual_distribution_plot_tmp.name, width=Inches(3))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = report.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Äußerung der Lehrperson (kurz)'
    hdr_cells[2].text = 'Code'
    for num, speech, code in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = speech
        row_cells[2].text = code

    report.add_paragraph('Anmerkung:')
    run = report.add_paragraph().add_run('IB = Invitation to expand ideas; IR = Invitation for reasoning; CH = Challenge; IC = Invitation for coordination; RD = Reflect on dialogue; NC = No code')
    font = run.font
    font.size = Pt(10)
    
    report.save(output_path)


def generate_report2(
    output_path: str, 
    group_name: str,
    num_pupils: int,
    num_participants: int,
    participation_rate: float,
    teacher_data: dict,
    student_data: dict,
    plot_distribution,  # matplotlib Figure
    num_impulses: int,
    plot_impulse_coding, # matplotlib Figure  
    impulse_table: pd.DataFrame
):
    # Neues Dokument
    doc = Document()

    # Schriften formatieren
    styles = doc.styles
    styles['Heading1'].element.rPr.rFonts.set(qn("w:asciiTheme"), "Aptos")
    styles['Heading1'].font.name = 'Aptos'
    styles['Heading1'].font.size = Pt(16)
    styles['Heading1'].font.bold = True
    styles['Heading1'].font.color.rgb = RGBColor(0, 0, 0)  # Schwarz
    styles['Heading1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles['Heading1'].paragraph_format.space_after = Pt(0)
    styles['Heading1'].paragraph_format.space_before = Pt(0)
    styles['Heading1'].paragraph_format.line_spacing = 1

    styles['Heading2'].font.name = 'Aptos'
    styles['Heading2'].font.size = Pt(12)
    styles['Heading2'].font.bold = True
    styles['Heading2'].font.color.rgb = RGBColor(0, 0, 0)  # Schwarz
    styles['Heading2'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles['Heading2'].paragraph_format.space_after = Pt(0)
    styles['Heading2'].paragraph_format.space_before = Pt(0)
    styles['Heading2'].paragraph_format.line_spacing = 1

    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(12)
    styles['Normal'].paragraph_format.space_after = Pt(0)
    styles['Normal'].paragraph_format.space_before = Pt(0)
    styles['Normal'].paragraph_format.line_spacing = 1

    # Ränder
    doc.sections[0].left_margin = Inches(0.5)
    doc.sections[0].right_margin = Inches(0.5)
    doc.sections[0].top_margin = Inches(0.5)
    doc.sections[0].bottom_margin = Inches(0.5)

    # === Titel ===
    doc.add_heading(f"Ergebnisse Simulationstraining – Gruppe {group_name}", level=1)

    # === Abschnitt: Quantitative Verteilung ===
    doc.add_heading("1. Quantitative Verteilung der Gesprächsanteile", level=2)
    doc.add_paragraph("").paragraph_format.line_spacing = 0.3
    doc.add_paragraph(f"Klassengröße: {num_pupils}\tAnzahl beteiligter Schüler:innen: {num_participants} (Beteiligungsquote: {participation_rate:.1f}%)")
    doc.add_paragraph("").paragraph_format.line_spacing = 0.5

    # === Tabelle: Gesprächsbeiträge ===
    par1 = doc.add_paragraph()
    par1.add_run("Tabelle ")
    par1.add_run("Gesprächsbeiträge der Lehrperson und der Schüler:innen").italic = True
    par1.paragraph_format.line_spacing = 1.2

    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Teilnehmende"
    hdr_cells[1].text = "Lehrperson"
#    hdr_cells[2].text = ""
 #   hdr_cells[3].text = ""
    hdr_cells[4].text = "Schüler:innen"
  #  hdr_cells[5].text = ""

    row1 = table.rows[1].cells
    row1[0].text = ""
    row1[1].text = "Anzahl"
    row1[2].text = "Länge in Wörtern"
    row1[3].text = ""
    row1[4].text = "Anzahl"
    row1[5].text = "Länge in Wörtern"

    row2 = table.rows[2].cells
    row2[0].text = ""
    row2[1].text = "N"
    row2[2].text = "M(SD)"
    row2[3].text = ""
    row2[4].text = "N"
    row2[5].text = "M(SD)"
    
    row3 = table.rows[3].cells
    row3[0].text = "Gesprächsbeiträge"
    row3[1].text = str(teacher_data["num"])
    row3[2].text = f"{str(teacher_data["words"])} ({str(teacher_data["mean_sd"])})"
    row3[3].text = ""
    row3[4].text = str(student_data["num"])
    row3[5].text = f"{str(student_data["words"])} ({str(student_data["mean_sd"])})"
    
    # Tabelle formatieren
    for cell in row2:
        cell.paragraphs[0].runs[0].italic = True   
    hdr_cells[1].merge(hdr_cells[2])
    hdr_cells[4].merge(hdr_cells[5])
    remove_table_borders(table)
    set_row_borders(hdr_cells, top=True, bottom=True)
    set_row_borders(row2[1:3], bottom=True)
    set_row_borders(row2[4:6], bottom=True)
    set_row_borders(row3, bottom=True)

    for row in table.rows:
        for cell in row.cells[2:6]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === Abbildung: Gesprächsverteilung als Plot ===
    doc.add_paragraph("")
    doc.add_paragraph("Abbildung")
    doc.add_paragraph().add_run("Gesprächsverteilung").italic = True
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpfile:
        plot_dis = plot_distribution
        plot_dis.figure.tight_layout()
        plot_dis.figure.set_size_inches(5.5, 2.9)
        plot_dis.figure.savefig(tmpfile.name, dpi=300, bbox_inches='tight')
        doc.add_picture(tmpfile.name)
    doc.add_paragraph("")

    # === Abschnitt: Qualitative Codierung ===
    doc.add_heading("2. Qualitative Codierung der Impulse der Lehrperson", level=2)
    doc.add_paragraph("").paragraph_format.line_spacing = 0.3
    doc.add_paragraph(f"Anzahl der Gesprächsimpulse: N = {num_impulses}")
    par2 = doc.add_paragraph()
    par2.add_run("Abbildung")
    par2.add_run(" Sprachliche Impulse der Lehrperson").italic = True

    # === Abbildung: Qualitative Verteilung als Plot ===
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpfile2:
        plot_qual = plot_impulse_coding
        plot_qual.figure.tight_layout()
        plot_qual.figure.set_size_inches(5.5, 2.9)
        plot_qual.figure.savefig(tmpfile2.name, dpi=300, bbox_inches='tight')
        doc.add_picture(tmpfile2.name)
    doc.add_paragraph("")

    # === Tabelle: Lehrerimpulse ===
    par3 = doc.add_paragraph()
    par3.add_run("Tabelle")
    par3.add_run(" Sprachliche Impulse der Lehrperson").italic = True

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Lehreräußerung (kurz)"
    hdr[2].text = "Code"

    for i, row in impulse_table.iterrows():
        row_cells = t.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = str(row['Lehreräußerung (kurz)'])
        row_cells[2].text = str(row['Shortcode'])

    # Tabellen-Header formatieren
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    # Schriftgröße für die Tabelle anpassen
    for row in t.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(8)

    # Breite der Zellen anpassen
    for row in t.rows:
        row.cells[0].width = Inches(0.3)
        row.cells[1].width = Inches(6.8)
        row.cells[2].width = Inches(0.5)

    # === Fußnote / Hinweis zu Codes ===
    doc.add_paragraph("")
    doc.add_paragraph(
        "Anmerkung: IB = Invitation to expand ideas; IR = Invitation for reasoning; CH = Challenge; "
        "IC = Invitation for coordination; RD = Reflect on dialogue; NC = No code"
    )

    # === Speichern ===
    doc.save(output_path)