from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Any

import docx
from bs4 import BeautifulSoup, NavigableString, Tag
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run

log = logging.getLogger(__name__)


class DocxStyle(Enum):
    """Enum for standard and custom .docx style names."""

    NORMAL = "Normal"
    HEADING_1 = "Heading 1"
    HEADING_2 = "Heading 2"
    HEADING_3 = "Heading 3"
    HEADING_4 = "Heading 4"
    HEADING_5 = "Heading 5"
    HEADING_6 = "Heading 6"
    LIST_NUMBER = "List Number"
    LIST_NUMBER_2 = "List Number 2"
    LIST_NUMBER_3 = "List Number 3"
    LIST_BULLET = "List Bullet"
    LIST_BULLET_2 = "List Bullet 2"
    LIST_BULLET_3 = "List Bullet 3"
    CODE = "Code"
    TABLE_GRID = "Table Grid"

    @property
    def type(self) -> WD_STYLE_TYPE:
        """Return the WD_STYLE_TYPE for the enum member."""
        if self.name.startswith("TABLE_"):
            return WD_STYLE_TYPE.TABLE
        # LIST_ styles are paragraph styles
        # CODE, HEADING, NORMAL are paragraph styles
        return WD_STYLE_TYPE.PARAGRAPH


@dataclass
class SimpleHtmlToDocx:
    """
    Convert restricted (simple, markdown-generated) HTML to .docx format.
    """

    list_indent: float = 0.4  # inches
    max_indent: float = 5.5  # inches
    max_recursion_depth: int = 100  # Prevent stack overflow on deeply nested documents

    def convert_html_string(self, html: str) -> Document:
        """
        Convert HTML string to Document object.
        """
        docx_template = Path(__file__).parent.resolve() / "templates" / "docx_template.docx"
        doc = docx.Document(str(docx_template))

        soup: BeautifulSoup = BeautifulSoup(html, "html.parser")

        self._process_element(soup, doc)
        return doc

    def convert_html_file(
        self, input_path: Path | str, output_path: Path | str | None = None
    ) -> None:
        """
        Convert HTML file to .docx file.
        """
        input_path = Path(input_path)
        try:
            html = input_path.read_text(encoding="utf-8")
            doc = self.convert_html_string(html)

            if output_path is None:
                output_path = input_path.with_suffix(".docx")

            doc.save(str(output_path))
        except Exception as e:
            raise RuntimeError(f"Error converting HTML file: {e}") from e

    def _process_element(
        self, element: Any, parent: Any, list_depth: int = 0, recursion_depth: int = 0
    ) -> None:
        """Recursively process HTML elements."""
        if recursion_depth > self.max_recursion_depth:
            return  # Prevent stack overflow

        if isinstance(element, NavigableString):
            text = str(element)
            # Collapse whitespace
            text = re.sub(r"\s+", " ", text)

            # Skip pure whitespace nodes between block elements
            if not text.strip():
                return

            if hasattr(parent, "add_run"):
                parent.add_run(text)
            return

        if not isinstance(element, Tag):
            return

        match element.name:
            case "h1" | "h2" | "h3" | "h4" | "h5" | "h6":
                level = int(element.name[1])
                p = parent.add_heading(level=level)
                self._process_children(element, p, list_depth, recursion_depth + 1)

            case "p":
                # If parent is already a paragraph, process children directly
                if isinstance(parent, Paragraph):
                    self._process_children(element, parent, list_depth, recursion_depth + 1)
                # Otherwise add a new paragraph to the parent
                elif hasattr(parent, "add_paragraph"):
                    p = parent.add_paragraph()
                    self._process_children(element, p, list_depth, recursion_depth + 1)
                # Fallback for unsupported parent
                else:
                    self._process_children(element, parent, list_depth, recursion_depth + 1)

            case "blockquote":
                # Process blockquote children directly to maintain paragraph structure
                for child in element.children:
                    if isinstance(child, Tag) and child.name == "p":
                        p = parent.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.right_indent = Inches(0.5)
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(10)
                        self._process_children(child, p, list_depth, recursion_depth + 1)
                    elif isinstance(child, NavigableString) and child.strip():
                        # Handle direct text in blockquote (not in a p tag)
                        p = parent.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.right_indent = Inches(0.5)
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(10)
                        p.add_run(child.strip())

            case "pre":
                p = parent.add_paragraph()
                p.style = self._get_style_name(self._get_document(parent), DocxStyle.CODE)
                text = element.get_text().strip()
                p.add_run(text)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)

            case "code":
                if hasattr(parent, "add_run"):
                    run = parent.add_run(element.get_text().strip())
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                elif hasattr(parent, "add_paragraph"):
                    p = parent.add_paragraph()
                    p.style = self._get_style_name(self._get_document(parent), DocxStyle.CODE)
                    p.add_run(element.get_text().strip())
                else:
                    # We're in an unknown container that can't add runs or paragraphs
                    # Just convert to text as a fallback
                    text = element.get_text().strip()
                    if hasattr(parent, "text"):
                        parent.text = text

            case "ul" | "ol":
                for li in element.find_all("li", recursive=False):
                    if isinstance(li, Tag):
                        self._process_list_item(
                            li,
                            self._get_document(parent),
                            element.name,
                            list_depth,
                            recursion_depth + 1,
                        )

            case "table":
                self._process_table(element, parent)

            case "hr":
                self._add_horizontal_rule(parent)

            case "br":
                if hasattr(parent, "add_run"):
                    parent.add_run().add_break()

            case "a" | "strong" | "b" | "em" | "i" | "code":
                if hasattr(parent, "add_run"):
                    # Preserve exact text without adding extra spaces
                    text = element.get_text()

                    # Apply appropriate formatting
                    run = parent.add_run(text)

                    if element.name in ["strong", "b"]:
                        run.font.bold = True
                    elif element.name in ["em", "i"]:
                        run.font.italic = True
                    elif element.name == "code":
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                    elif element.name == "a":
                        run.font.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)

                        # Handle href for links
                        href = element.get("href", "")
                        if isinstance(href, list):
                            href = " ".join(href)
                        href = str(href)
                        if href:
                            self._add_hyperlink(parent, run, href, text)

            case "img":
                # Support for image would go here if needed
                # Would require additional code to download images and add them to the document
                pass

            case _:
                # For any other tags, just process children
                self._process_children(element, parent, list_depth, recursion_depth + 1)

    def _process_children(
        self, element: Tag, parent: Any, list_depth: int = 0, recursion_depth: int = 0
    ) -> None:
        """Process all children of an element."""
        # Define inline elements that should preserve surrounding whitespace
        inline_tags = {"a", "strong", "b", "em", "i", "code", "span", "u", "s", "sub", "sup"}

        children = list(element.children)

        for i, child in enumerate(children):
            if isinstance(parent, Run) and isinstance(child, NavigableString):
                parent.text += str(child)
            elif isinstance(parent, Run) and isinstance(child, Tag):
                self._process_element(child, parent, list_depth, recursion_depth + 1)
            else:
                if isinstance(child, NavigableString) and isinstance(parent, Paragraph):
                    text = str(child)
                    # Collapse whitespace
                    text = re.sub(r"\s+", " ", text)

                    # Skip pure whitespace
                    if not text.strip():
                        continue

                    # Trim leading whitespace if at start of paragraph
                    if i == 0 or all(
                        isinstance(c, NavigableString) and not c.strip() for c in children[:i]
                    ):
                        text = text.lstrip()

                    # Trim trailing whitespace if at end or before non-inline element
                    should_trim_trailing = False
                    if i == len(children) - 1:
                        should_trim_trailing = True
                    else:
                        # Look for the next non-whitespace sibling
                        for j in range(i + 1, len(children)):
                            next_child = children[j]
                            if isinstance(next_child, NavigableString):
                                if next_child.strip():  # Found non-whitespace text
                                    break
                            elif isinstance(next_child, Tag):
                                # Trim if next element is not inline
                                should_trim_trailing = next_child.name not in inline_tags
                                break

                    if should_trim_trailing:
                        text = text.rstrip()

                    if text:
                        parent.add_run(text)
                else:
                    self._process_element(child, parent, list_depth, recursion_depth + 1)

    def _process_list_item(
        self, li: Tag, doc: Document, list_type: str, depth: int, recursion_depth: int = 0
    ) -> None:
        """Process a list item by applying pre-defined list styles from the template."""

        style_enum: DocxStyle
        if list_type == "ol":
            if depth == 0:
                style_enum = DocxStyle.LIST_NUMBER
            elif depth == 1:
                style_enum = DocxStyle.LIST_NUMBER_2
            else:
                style_enum = DocxStyle.LIST_NUMBER_3
        else:  # ul
            if depth == 0:
                style_enum = DocxStyle.LIST_BULLET
            elif depth == 1:
                style_enum = DocxStyle.LIST_BULLET_2
            else:
                style_enum = DocxStyle.LIST_BULLET_3

        try:
            # Attempt to apply the style from the template.
            style_name_str = self._get_style_name(doc, style_enum)
            p = doc.add_paragraph(style=style_name_str)
        except KeyError:
            # Fallback if style is not in the template (will likely not look like a list)
            log.warning(
                f"List style '{style_enum.value}' not found in template. Adding paragraph without list styling."
            )
            p = doc.add_paragraph()
            # Apply manual indentation as a minimal fallback
            p.paragraph_format.left_indent = Inches(0.5 + depth * 0.5)
            if depth > 0:  # Hanging indent for nested, simple indent for top level
                p.paragraph_format.first_line_indent = Inches(-0.25)

        # Process direct content (excluding nested lists) into this paragraph
        direct_content = []
        nested_lists = []

        # Separate direct content from nested lists
        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                nested_lists.append(child)
            else:
                direct_content.append(child)

        # Find the last non-empty text node if we have nested lists
        last_text_index = -1
        if nested_lists:
            for i in range(len(direct_content) - 1, -1, -1):
                content = direct_content[i]
                if isinstance(content, NavigableString) and content.strip():
                    last_text_index = i
                    break

        # Process direct content
        for i, content in enumerate(direct_content):
            if isinstance(content, NavigableString) and i == last_text_index:
                # This is the last non-empty text before nested lists
                text = str(content)
                text = re.sub(r"\s+", " ", text)
                text = text.strip()  # Trim both sides for the last text
                if text:
                    p.add_run(text)
            else:
                self._process_element(content, p, depth, recursion_depth + 1)

        # Process nested lists (each gets its own paragraph with proper numbering)
        for nested_list in nested_lists:
            for nested_li in nested_list.find_all("li", recursive=False):
                if isinstance(nested_li, Tag):
                    self._process_list_item(
                        nested_li, doc, nested_list.name, depth + 1, recursion_depth + 1
                    )

    def _process_table(self, table: Tag, doc: Document) -> None:
        """Process a table element."""
        rows_elements = table.find_all("tr")
        if not rows_elements:
            return

        # Filter for actual Tag objects for rows
        rows_tags = [r for r in rows_elements if isinstance(r, Tag)]
        if not rows_tags:
            return

        # Calculate dimensions
        max_cols = max(len(row.find_all(["td", "th"])) for row in rows_tags)

        # Create table
        docx_table = doc.add_table(rows=len(rows_tags), cols=max_cols if max_cols > 0 else 1)
        docx_table.style = self._get_style_name(doc, DocxStyle.TABLE_GRID)  # Add borders

        # Fill cells
        for row_idx, row_element in enumerate(rows_tags):
            cell_elements = row_element.find_all(["td", "th"])
            for col_idx, cell_element in enumerate(cell_elements):
                if not isinstance(cell_element, Tag):
                    continue

                docx_cell = docx_table.cell(row_idx, col_idx)
                docx_cell.text = ""  # Clear default paragraph
                p = docx_cell.add_paragraph()

                # Bold for header cells
                if cell_element.name == "th":
                    run = p.add_run("")
                    run.font.bold = True

                # Handle alignment
                self._apply_cell_alignment(cell_element, p)

                # Process cell content
                self._process_children(cell_element, p)

    def _apply_cell_alignment(self, cell_element: Tag, paragraph: Paragraph) -> None:
        """Extract and apply alignment from cell attributes."""
        # Check both align attribute and style attribute for alignment info
        align_attr = cell_element.get("align", "") or ""
        style_attr = cell_element.get("style", "") or ""

        # Convert to string if it's a list
        if isinstance(align_attr, list):
            align_attr = " ".join(align_attr)
        if isinstance(style_attr, list):
            style_attr = " ".join(style_attr)

        # Extract text-align from style attribute if present
        text_align = ""
        if style_attr:
            align_match = re.search(r"text-align:\s*(\w+)", style_attr)
            if align_match:
                text_align = align_match.group(1).lower()

        # Determine alignment
        if "center" in align_attr or text_align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "right" in align_attr or text_align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # else: leave at default (left alignment)

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Add a horizontal rule to the document."""
        p = doc.add_paragraph("* * *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def _add_hyperlink(self, paragraph: Paragraph, run: Any, url: str, text: str) -> None:
        """Add hyperlink to the document."""
        # Create a new run with the same text and formatting
        new_run = paragraph.add_run(text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if hasattr(run, "font") and hasattr(run.font, "color") and hasattr(run.font.color, "rgb"):
            new_run.font.color.rgb = run.font.color.rgb

        # Remove the original run
        run_parent = run._element.getparent()
        run_parent.remove(run._element)

        # Add hyperlink relationship
        r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create hyperlink and add the new run's element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        hyperlink.append(new_run._element)

        # Add hyperlink to paragraph
        paragraph._p.append(hyperlink)

    def _get_document(self, parent: Any) -> Document:
        """Get the Document object from any parent (Document, Paragraph, etc.)."""
        if isinstance(parent, Document):
            return parent
        elif hasattr(parent, "_part") and hasattr(parent._part, "document"):
            return parent._part.document
        else:
            raise ValueError(f"Cannot get document from parent of type {type(parent)}")

    def _get_style_name(self, doc: Document, style_enum_member: DocxStyle) -> str:
        """
        Ensures a style exists by its DocxStyle enum member, returning its string name.
        If a style is missing, it logs a warning and attempts to create a basic version.
        It's strongly recommended the template defines these styles properly.
        """
        style_name = style_enum_member.value
        style_type = style_enum_member.type

        try:
            style = doc.styles[style_name]
            if style.type != style_type:
                log.warning(
                    f"Style '{style_name}' exists in template but is type '{style.type}', not expected '{style_type}'."
                    f" Formatting may be incorrect."
                )
            return style_name
        except KeyError:
            log.warning(
                f"Style '{style_name}' not found in template. "
                f"Attempting to add a basic '{style_type.name}' style. "
                f"For proper formatting, define this style in your .docx template."
            )
            doc.styles.add_style(style_name, style_type)
            created_style = doc.styles[style_name]

            if style_type == WD_STYLE_TYPE.PARAGRAPH:
                if isinstance(created_style, ParagraphStyle):
                    if style_enum_member == DocxStyle.CODE:
                        created_style.font.name = "Courier New"
                        created_style.font.size = Pt(9)
                        created_style.paragraph_format.space_before = Pt(6)
                        created_style.paragraph_format.space_after = Pt(6)
                    elif style_name.startswith("List"):
                        created_style.paragraph_format.left_indent = Inches(0.5)
                else:
                    log.error(
                        f"Style '{style_name}' was added but could not be fetched as ParagraphStyle."
                    )

            return style_name


## Tests


_SAMPLE_HTML = """
<h1>Document Title</h1>
<p>This is a <strong>bold</strong> paragraph with <em>italic</em> text and a <a href="https://example.com">link</a>.</p>
<h2>Lists</h2>
<ul>
    <li>Unordered item 1</li>
    <li>Unordered item 2
        <ul>
            <li>Nested item A</li>
            <li>Nested item B</li>
        </ul>
    </li>
</ul>
<ol>
    <li>Ordered item 1</li>
    <li>Ordered item 2</li>
</ol>
<h2>Code</h2>
<pre><code>def hello_world():
    print("Hello, World!")
</code></pre>
<p>Inline <code>code</code> example</p>
<h2>Blockquote</h2>
<blockquote>
    <p>This is a blockquote with <strong>formatting</strong>.</p>
    <p>And a second paragraph.</p>
</blockquote>
<hr>
<h2>Table</h2>
<table>
    <tr>
        <th>Header 1</th>
        <th>Header 2</th>
        <th style="text-align:right">Right Aligned</th>
    </tr>
    <tr>
        <td>Row 1, Col 1</td>
        <td>Row 1, Col 2</td>
        <td align="right">1234</td>
    </tr>
    <tr>
        <td>Row 2, Col 1</td>
        <td>Row 2, Col 2</td>
        <td align="right">5678</td>
    </tr>
</table>
"""

# FIXME: Not roundtripping nested lists correctly as we lose the information we need.
# Fix code and then nested items below to be indented.

_EXPECTED_HTML = (
    """<h1>Document Title</h1>"""
    """<p>This is a <strong>bold</strong> paragraph with <em>italic</em> text and a <a href="https://example.com">link</a>.</p>"""
    """<h2>Lists</h2>"""
    """<ul><li>Unordered item 1</li><li>Unordered item 2</li><li>Nested item A</li><li>Nested item B</li></ul>"""
    """<ol><li>Ordered item 1</li><li>Ordered item 2</li></ol>"""
    """<h2>Code</h2>"""
    """<p>def hello_world():<br />    print(&quot;Hello, World!&quot;)</p>"""
    """<p>Inline code example</p>"""
    """<h2>Blockquote</h2>"""
    """<p>This is a blockquote with <strong>formatting</strong>.</p><p>And a second paragraph.</p>"""
    """<p>* * *</p>"""
    """<h2>Table</h2>"""
    """<table><tr><td><p>Header 1</p></td><td><p>Header 2</p></td><td><p>Right Aligned</p></td></tr><tr><td><p>Row 1, Col 1</p></td><td><p>Row 1, Col 2</p></td><td><p>1234</p></td></tr><tr><td><p>Row 2, Col 1</p></td><td><p>Row 2, Col 2</p></td><td><p>5678</p></td></tr></table>"""
)

_EXPECTED_MD = r"""
# Document Title

This is a **bold** paragraph with *italic* text and a [link](https://example.com).

## Lists

* Unordered item 1
* Unordered item 2
* Nested item A
* Nested item B
1. Ordered item 1
2. Ordered item 2

## Code

def hello\_world():\
 print("Hello, World!")

Inline code example

## Blockquote

This is a blockquote with **formatting**.

And a second paragraph.

\* \* \*

## Table

| Header 1 | Header 2 | Right Aligned |
| --- | --- | --- |
| Row 1, Col 1 | Row 1, Col 2 | 1234 |
| Row 2, Col 1 | Row 2, Col 2 | 5678 |
"""


# Normalize for cleaner testing
def _normalize_html(html_str: str) -> str:
    soup = BeautifulSoup(html_str, "html.parser")
    return str(soup.prettify(formatter="html5")).strip()


def test_html_to_docx_conversion():
    from kash.kits.docs.doc_formats import markitdown_convert

    converter = SimpleHtmlToDocx()

    # Test string conversion
    doc = converter.convert_html_string(_SAMPLE_HTML)
    assert doc is not None

    # Verify document has content
    assert len(doc.paragraphs) > 0
    assert len(doc.tables) > 0

    # Test file conversion
    out_dir = Path("tmp")
    out_dir.mkdir(exist_ok=True)
    temp_html_path = out_dir / "test_doc.html"
    temp_docx_path = out_dir / "test_doc.docx"

    with open(temp_html_path, "w") as temp_html:
        temp_html.write(_SAMPLE_HTML)

    converter.convert_html_file(temp_html_path, temp_docx_path)

    # Verify file exists and has content
    assert temp_docx_path.exists()
    assert temp_docx_path.stat().st_size > 0

    md = markitdown_convert.docx_to_md(temp_docx_path)

    # Normalize HTML for comparison
    assert md.raw_html
    actual_html_normalized = _normalize_html(md.raw_html)
    expected_html_normalized = _normalize_html(_EXPECTED_HTML)

    # Print for debugging if needed
    if actual_html_normalized != expected_html_normalized:
        import difflib

        print("=== HTML DIFF ===")
        diff = difflib.unified_diff(
            expected_html_normalized.splitlines(keepends=True),
            actual_html_normalized.splitlines(keepends=True),
            fromfile="expected",
            tofile="actual",
            lineterm="",
        )
        print("".join(diff))

    print("=== DOCX OUTPUT ===")
    print(f"{temp_docx_path}")

    # Compare normalized HTML and raw markdown
    assert actual_html_normalized == expected_html_normalized
    assert md.markdown.strip() == _EXPECTED_MD.strip()

    # Clean up temp files
    # if os.path.exists(temp_html_path):
    #     os.unlink(temp_html_path)
    # if os.path.exists(temp_docx_path):
    #     os.unlink(temp_docx_path)
