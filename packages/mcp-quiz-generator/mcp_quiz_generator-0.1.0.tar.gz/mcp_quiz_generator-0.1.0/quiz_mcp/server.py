#!/usr/bin/env python3
"""
MCP Quiz Generator Server

A Model Context Protocol server for generating HTML and Word quiz files 
from Markdown content with support for multiple question types.
"""

import os
import re
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import click
import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from jinja2 import Environment, FileSystemLoader, select_autoescape
from mcp.server.fastmcp import FastMCP

# Default configuration
DEFAULT_OUTPUT_FOLDER = "data"

# Create FastMCP server instance
mcp = FastMCP("mcp-quiz-generator")

class QuizGenerator:
    """Quiz generator utility class."""

    def __init__(self, output_folder: str = DEFAULT_OUTPUT_FOLDER):
        """Initialize the quiz generator.
        
        Args:
            output_folder: Directory to save generated files
        """
        self.output_folder = Path(output_folder)
        
        # Ensure output directory exists
        self.output_folder.mkdir(exist_ok=True)
        
        # Setup template environment
        template_dir = Path(__file__).parent / "templates"
        if not template_dir.exists():
            # Fallback to app directory for backward compatibility
            template_dir = Path(__file__).parent / "app" / "static"
        
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=select_autoescape(['html', 'xml'])
        )

    def markdown_to_html(self, markdown_content: str) -> Tuple[str, str]:
        """Convert markdown content to HTML and return HTML content and filename."""
        # Use the same extensions as the original code
        # extensions = [
        #     "tables", 
        #     "mcp_quiz_generator.extensions.checkbox", 
        #     "mcp_quiz_generator.extensions.radio",
        #     "mcp_quiz_generator.extensions.textbox"
        # ]
        extensions = [
            "tables", 
            "quiz_mcp.app.extensions.checkbox",
            "quiz_mcp.app.extensions.radio", 
            "quiz_mcp.app.extensions.textbox"
        ]
        
        # Convert markdown to HTML
        html = markdown.markdown(
            markdown_content,
            extensions=extensions,
            output_format="html5",
            extension_configs={}
        )
        
        # Try to render JavaScript template
        try:
            javascript = self.jinja_env.get_template('app.js').render()
        except:
            # Fallback JavaScript if template not found
            javascript = """
            // Basic quiz functionality
            document.addEventListener('DOMContentLoaded', function() {
                console.log('Quiz loaded');
            });
            """
        
        # Try to render base HTML template
        try:
            test_html = self.jinja_env.get_template('base.html').render(
                content=html,
                javascript=javascript
            )
        except:
            # Fallback HTML structure
            test_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Quiz</title>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    .question {{ margin: 20px 0; }}
                    .options {{ margin-left: 20px; }}
                </style>
            </head>
            <body>
                {html}
                <script>{javascript}</script>
            </body>
            </html>
            """
        
        # Try to render wrapper template
        try:
            final_html = self.jinja_env.get_template('wrapper.html').render(content=test_html)
        except:
            # Use test_html as final if wrapper not found
            final_html = test_html
        
        # Generate timestamped filename
        filename = f"{int(time.time())}.html"
        
        return final_html, filename

    def create_word_document(self, markdown_content: str, include_answers: bool = True) -> Document:
        """Create Word document with optional answer inclusion."""
        doc = Document()
        
        # Add title
        if include_answers:
            title = doc.add_heading('Quiz Document (Answer Key)', 0)
        else:
            title = doc.add_heading('Quiz Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content line by line
        lines = markdown_content.split('\n')
        question_counter = 0
        option_counter = 0
        
        for line in lines:
            if not line.strip():
                continue
                
            # Handle headers
            if line.strip().startswith('#'):
                level = len(line.strip()) - len(line.strip().lstrip('#'))
                title_text = line.strip().lstrip('#').strip()
                if level == 1:
                    doc.add_heading(title_text, 1)
                else:
                    doc.add_heading(title_text, 2)
            
            # Handle numbered questions
            elif re.match(r'^\d+\.', line.strip()):
                question_counter += 1
                option_counter = 0
                question_text = re.sub(r'^\d+\.\s*', '', line.strip())
                question_para = doc.add_paragraph()
                question_para.add_run(f"{question_counter}. ").bold = True
                question_para.add_run(question_text)
            
            # Handle option lines (starting with 4 spaces and -)
            elif line.startswith('    -'):
                option_text = line[6:].strip()  # Remove '    - '
                
                # Handle radio button options (x) or ( )
                if option_text.startswith('(x)') or option_text.startswith('( )'):
                    option_counter += 1
                    clean_text = option_text[3:].strip()
                    is_correct = option_text.startswith('(x)')
                    
                    option_para = doc.add_paragraph()
                    option_letter = chr(ord('A') + option_counter - 1)
                    
                    if include_answers and is_correct:
                        option_para.add_run(f"{option_letter}. ").bold = True
                        option_para.add_run(clean_text).bold = True
                        option_para.add_run(" ✓").bold = True
                    else:
                        option_para.add_run(f"{option_letter}. ")
                        option_para.add_run(clean_text)
                
                # Handle checkbox options [x] or [ ]
                elif option_text.startswith('[x]') or option_text.startswith('[ ]'):
                    option_counter += 1
                    clean_text = option_text[3:].strip()
                    is_correct = option_text.startswith('[x]')
                    
                    option_para = doc.add_paragraph()
                    option_letter = chr(ord('A') + option_counter - 1)
                    
                    if include_answers and is_correct:
                        option_para.add_run(f"{option_letter}. ").bold = True
                        option_para.add_run(clean_text).bold = True
                        option_para.add_run(" ✓").bold = True
                    else:
                        option_para.add_run(f"{option_letter}. ")
                        option_para.add_run(clean_text)
                
                # Handle text input questions R:=
                elif option_text.startswith('R:=') or option_text.startswith('r:='):
                    answer_text = option_text[3:].strip()
                    option_para = doc.add_paragraph()
                    if include_answers:
                        option_para.add_run("Answer: ").bold = True
                        option_para.add_run(f"_______ (Correct answer: {answer_text})")
                    else:
                        option_para.add_run("Answer: ")
                        option_para.add_run("_______________________")
                
                # Other option formats
                else:
                    option_para = doc.add_paragraph(style='List Bullet')
                    option_para.add_run("• ")
                    option_para.add_run(option_text)
            
            # Handle horizontal lines
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            
            # Handle regular paragraphs
            else:
                if line:
                    doc.add_paragraph(line)
        
        # Add instructions page
        doc.add_page_break()
        if include_answers:
            instructions_heading = doc.add_heading('Answer Key Instructions', 1)
            instructions = [
                "• This is the answer key version",
                "• Correct answers are marked with ✓ and shown in bold",
                "• For multiple choice questions, select the option marked with ✓",
                "• For multiple answer questions, select all options marked with ✓",
                "• For text input questions, use the provided correct answer"
            ]
        else:
            instructions_heading = doc.add_heading('Instructions', 1)
            instructions = [
                "• For multiple choice questions, select one correct answer (A, B, C, or D)",
                "• For multiple answer questions, select all correct answers",
                "• For text input questions, write your answer in the blank space",
                "• Please write clearly and legibly"
            ]
        
        for instruction in instructions:
            doc.add_paragraph(instruction, style='List Bullet')
        
        return doc

# Global quiz generator instance
# quiz_generator = QuizGenerator()

@mcp.tool()
def generate_quiz(markdown_content: str, format_type: str = "both", custom_filename: str = "") -> str:
    """Generate quiz files from markdown content in specified format(s).
    
    This tool converts markdown quiz content to HTML and/or Word documents based on the format_type parameter.
    
    Markdown Text Format Requirements:
    The markdown content must follow this specific format for different question types:
    
    #### Single Choice Questions
    1. What are the test automation frameworks developed by MaxSoft?
        - [x] IntelliAPI  (correct answer marked with [x])
        - [ ] WebBot      (incorrect answers marked with [ ])
        - [ ] Gauge
        - [ ] Selenium

    #### Multiple Choice Questions
    2. What are the test automation frameworks developed by MaxSoft?
        - [x] IntelliAPI  (multiple correct answers marked with [x])
        - [x] WebBot      (multiple correct answers marked with [x])
        - [ ] Gauge       (incorrect answers marked with [ ])
        - [ ] Selenium

    #### True/False Questions
    3. MaxSoft is a software company.
        - (x) True        (correct answer marked with (x))
        - ( ) False       (incorrect answer marked with ( ))

    4. The domain of MaxSoft is test automation framework development.
        - (x) True
        - ( ) False

    #### Short Answer Questions
    5. Who is the Co-Founder of MaxSoft?
        - R:= Osanda      (answer marked with R:= followed by the correct answer)
    
    Args:
        markdown_content: Markdown content following the specified format requirements above
        format_type: Output format - "html" for HTML only, "word" for Word only, "both" for both formats (default: "word")
        custom_filename: Optional custom filename (without extension). If not provided, timestamp will be used.
    
    Returns:
        Success message with file paths for the generated format(s)
    """
    if not markdown_content.strip():
        return "Error: markdown_content cannot be empty"
    
    if format_type not in ["html", "word", "both"]:
        return "Error: format_type must be 'html', 'word', or 'both'"
    
    try:
        timestamp = int(time.time())
        
        if custom_filename.strip():
            filename_prefix = custom_filename.strip()
        else:
            filename_prefix = str(timestamp)
        
        results = []
        
        # Generate HTML if requested
        if format_type in ["html", "both"]:
            html_content, _ = quiz_generator.markdown_to_html(markdown_content)
            html_filename = f"{filename_prefix}.html"
            html_file_path = quiz_generator.output_folder / html_filename
            html_file_path.write_text(html_content, encoding='utf-8')
            
            results.append(f"HTML (Interactive):\n  File: {html_file_path}")
        
        # Generate Word if requested
        if format_type in ["word", "both"]:
            # Generate answer key version
            answer_key_doc = quiz_generator.create_word_document(markdown_content, include_answers=True)
            answer_key_filename = f"{filename_prefix}_answer_key.docx"
            answer_key_path = quiz_generator.output_folder / answer_key_filename
            answer_key_doc.save(str(answer_key_path))
            
            # Generate quiz version (without answers)
            quiz_doc = quiz_generator.create_word_document(markdown_content, include_answers=False)
            quiz_filename = f"{filename_prefix}_quiz.docx"
            quiz_path = quiz_generator.output_folder / quiz_filename
            quiz_doc.save(str(quiz_path))
            
            results.append(f"Word Answer Key:\n  File: {answer_key_path}")
            results.append(f"Word Quiz (for students):\n  File: {quiz_path}")
        
        format_name = {"html": "HTML", "word": "Word", "both": "All formats"}[format_type]
        return f"{format_name} generated successfully!\n\n" + "\n\n".join(results)

    except Exception as e:
        return f"Error generating quiz files: {str(e)}"

@click.command()
@click.option(
    "--output-folder",
    default=DEFAULT_OUTPUT_FOLDER,
    help=f"Output folder for generated files (default: {DEFAULT_OUTPUT_FOLDER})"
)
def main(output_folder: str) -> None:
    """Run the MCP Quiz Generator server."""
    print(f"MCP Quiz Generator Server starting...")
    print(f"Output folder: {output_folder}")
    
    # Update global quiz generator with custom settings
    global quiz_generator
    quiz_generator = QuizGenerator(output_folder=output_folder)
    
    # Run the FastMCP server
    mcp.run()
    # mcp.run(transport = "streamable-http")

if __name__ == "__main__":
    main()
