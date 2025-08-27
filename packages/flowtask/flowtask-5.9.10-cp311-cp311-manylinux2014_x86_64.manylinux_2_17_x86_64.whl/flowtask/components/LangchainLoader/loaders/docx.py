from typing import List
from pathlib import PurePath
import mammoth
import docx
from markdownify import markdownify as md
from langchain.docstore.document import Document
from .abstract import AbstractLoader


class MSWordLoader(AbstractLoader):
    """
    Load Microsoft Docx as Langchain Documents.
    """
    def extract_text(self, path):
        """Extract text from a docx file.

        Args:
            path (Path): The source of the data.

        Returns:
            str: The extracted text.
        """
        doc = docx.Document(str(path))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    async def _load_document(self, path: PurePath) -> List[Document]:
        """Load data from a source and return it as a Langchain Document.

        Args:
            path (Path): The source of the data.

        Returns:
            List[Document]: A list of Langchain Documents.
        """
        self.logger.info(
            f"Loading Word file: {path}"
        )
        docs = []
        with open(path, "rb") as docx_file:
            doc = docx.Document(str(path))
            properties = doc.core_properties
            result = mammoth.convert_to_html(docx_file)
            # text_file = mammoth.extract_raw_text(docx_file) # Use text File for summary
            html = result.value  # The generated HTML
            md_text = md(html)  # The generated Markdown

            print(f"Type of HTML result: {type(html)}")
            print(f"Length of HTML: {len(html)}")
            print(f"First 100 characters: {html[:100]}")
            print(f"Messages from conversion: {result.messages}")
            # TODO: add summarization and translation if requested
            summary = ''
            # try:
            #     summary = self.get_summary_from_text(md_text, use_gpu=True)
            # except Exception:
            #     summary = ''
            document_meta = {
                "author": properties.author,
                "version": properties.version,
                "title": properties.title,
                # "created": properties.created.strftime("%Y-%m-%d %H:%M:%S"),
                # "last_modified": properties.modified.strftime("%Y-%m-%d %H:%M:%S")
            }
            metadata = self.create_metadata(
                path=path,
                doctype=self.doctype,
                source_type=self._source_type,
                summary=summary,
                doc_metadata=document_meta
            )
            # Create document-level context
            document_context = f"File Name: {path.name}\n"
            document_context += f"Document Type: {self.doctype}\n"
            document_context += f"Source Type: {self._source_type}\n"
            document_context += f"Summary: {summary}\n"
            document_context += "======\n"
            # splitting the content:
            for chunk in self.markdown_splitter.split_text(md_text):
                _idx = {
                    **metadata
                }
                docs.append(
                    Document(
                        page_content=document_context + chunk,
                        metadata=_idx
                    )
                )
        return docs
