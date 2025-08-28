from io import BytesIO

import emoji
from docx import Document
from osd_text_extractor.domain.interfaces import TextExtractor
from osd_text_extractor.infrastructure.exceptions import ExtractionError


class DOCXExtractor(TextExtractor):
    @staticmethod
    def extract_plain_text(file_content: bytes) -> str:
        try:
            doc = Document(BytesIO(file_content))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            text = "\n".join(text_parts)
            return emoji.replace_emoji(text, replace=" ")
        except Exception as e:
            raise ExtractionError("Failed to extract DOCX text") from e
