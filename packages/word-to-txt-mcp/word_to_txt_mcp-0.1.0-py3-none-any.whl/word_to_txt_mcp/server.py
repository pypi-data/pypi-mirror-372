"""FastMCP Word文档转文本分析服务器核心模块"""

import os
import tempfile
from docx import Document
from fastmcp import FastMCP


def convert_word_to_text(word_file_path):
    """
    将Word文档转换为文本内容
    
    参数:
        word_file_path (str): Word文档的文件路径
    
    返回值:
        str: 提取的文本内容
    
    异常:
        FileNotFoundError: 当Word文件不存在时抛出
        Exception: 当转换过程中出现错误时抛出
    """
    try:
        # 检查Word文件是否存在
        if not os.path.exists(word_file_path):
            raise FileNotFoundError(f"Word文件不存在: {word_file_path}")
        
        # 读取Word文档
        doc = Document(word_file_path)
        
        # 提取所有段落的文本
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 只添加非空段落
                text_content.append(paragraph.text.strip())
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n'.join(text_content)
        
    except Exception as e:
        raise Exception(f"转换Word文档时出错: {str(e)}")


def process_word_document_func(file_path: str) -> str:
    """
    处理Word文档，将其转换为文本并进行基础分析
    
    参数:
        file_path (str): Word文档的文件路径
    
    返回值:
        str: 包含文档内容和基础分析的结果
    """
    try:
        # 转换Word文档为文本
        text_content = convert_word_to_text(file_path)
        
        if not text_content.strip():
            return "文档内容为空或无法提取文本内容。"
        
        # 基础文档分析
        word_count = len(text_content.split())
        char_count = len(text_content)
        line_count = len(text_content.split('\n'))
        
        # 构建分析结果
        analysis_result = f"""=== Word文档分析结果 ===

文档统计信息:
- 字符数: {char_count}
- 单词数: {word_count}
- 行数: {line_count}

=== 文档内容 ===
{text_content}

=== 分析完成 ===
文档已成功转换为文本格式，可以进行进一步的AI分析处理。"""
        
        return analysis_result
        
    except Exception as e:
        return f"处理文档时出错: {str(e)}"


def analyze_document_content_func(text_content: str, analysis_type: str = "summary") -> str:
    """
    分析文档内容
    
    参数:
        text_content (str): 要分析的文本内容
        analysis_type (str): 分析类型 (summary, keywords, structure)
    
    返回值:
        str: 分析结果
    """
    try:
        if not text_content.strip():
            return "文本内容为空，无法进行分析。"
        
        # 基础统计
        words = text_content.split()
        sentences = text_content.split('。')
        paragraphs = [p.strip() for p in text_content.split('\n') if p.strip()]
        
        if analysis_type == "summary":
            result = f"""=== 文档摘要分析 ===

基础统计:
- 总字符数: {len(text_content)}
- 总词数: {len(words)}
- 句子数: {len([s for s in sentences if s.strip()])}
- 段落数: {len(paragraphs)}

内容预览:
{text_content[:500]}{'...' if len(text_content) > 500 else ''}

=== 分析建议 ===
此文档已转换为纯文本格式，建议使用专业的AI模型进行深度内容分析、摘要生成或关键信息提取。"""
            
        elif analysis_type == "keywords":
            # 简单的关键词提取（基于词频）
            word_freq = {}
            for word in words:
                clean_word = word.strip('.,!?;:"()[]{}').lower()
                if len(clean_word) > 2:  # 只统计长度大于2的词
                    word_freq[clean_word] = word_freq.get(clean_word, 0) + 1
            
            # 获取频率最高的10个词
            top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
            
            result = f"""=== 关键词分析 ===

高频词汇 (前10个):
"""
            for i, (word, freq) in enumerate(top_words, 1):
                result += f"{i}. {word} (出现{freq}次)\n"
            
        elif analysis_type == "structure":
            result = f"""=== 文档结构分析 ===

文档结构:
- 总段落数: {len(paragraphs)}
- 平均段落长度: {sum(len(p) for p in paragraphs) // len(paragraphs) if paragraphs else 0} 字符

段落预览:
"""
            for i, para in enumerate(paragraphs[:5], 1):
                result += f"\n段落 {i}: {para[:100]}{'...' if len(para) > 100 else ''}"
            
            if len(paragraphs) > 5:
                result += f"\n\n... 还有 {len(paragraphs) - 5} 个段落"
        
        else:
            result = "不支持的分析类型。支持的类型: summary, keywords, structure"
        
        return result
        
    except Exception as e:
        return f"分析文档内容时出错: {str(e)}"


def create_mcp_server(server_name: str = "Document Analysis Server"):
    """
    创建并配置MCP服务器实例
    
    参数:
        server_name (str): 服务器名称
    
    返回值:
        FastMCP: 配置好的MCP服务器实例
    """
    # 创建服务器
    mcp = FastMCP(server_name)
    
    @mcp.tool
    def process_word_document(file_path: str) -> str:
        """处理Word文档工具"""
        return process_word_document_func(file_path)
    
    @mcp.tool
    def analyze_document_content(text_content: str, analysis_type: str = "summary") -> str:
        """分析文档内容工具"""
        return analyze_document_content_func(text_content, analysis_type)
    
    @mcp.tool
    def echo_tool(text: str) -> str:
        """Echo工具"""
        return text
    
    @mcp.resource("document://help")
    def document_help_resource() -> str:
        """
        提供文档处理功能的帮助信息
        
        返回值:
            str: 帮助信息
        """
        return """=== 文档分析服务帮助 ===

可用工具:
1. process_word_document(file_path) - 处理Word文档并转换为文本
2. analyze_document_content(text_content, analysis_type) - 分析文档内容
   - analysis_type可选: summary, keywords, structure
3. echo_tool(text) - 回显文本

使用流程:
1. 上传Word文档
2. 使用process_word_document处理文档
3. 使用analyze_document_content进行深度分析
4. 将结果提供给AI模型进行进一步处理

支持的文档格式: .docx, .doc"""
    
    @mcp.resource("document://status/{file_path}")
    def document_status_resource(file_path: str) -> str:
        """
        检查文档状态
        
        参数:
            file_path (str): 文档文件路径
        
        返回值:
            str: 文档状态信息
        """
        try:
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                file_ext = os.path.splitext(file_path)[1].lower()
                
                status = f"""文档状态:
- 文件路径: {file_path}
- 文件大小: {file_size} 字节
- 文件类型: {file_ext}
- 状态: 文件存在
- 支持处理: {'是' if file_ext in ['.docx', '.doc'] else '否'}"""
                return status
            else:
                return f"文档状态: 文件不存在 - {file_path}"
        except Exception as e:
            return f"检查文档状态时出错: {str(e)}"
    
    @mcp.prompt("analyze_document")
    def analyze_document_prompt(file_path: str, analysis_requirements: str = "基础分析") -> str:
        """
        生成文档分析提示
        
        参数:
            file_path (str): 文档文件路径
            analysis_requirements (str): 分析需求
        
        返回值:
            str: 分析提示
        """
        return f"""请分析以下Word文档:

文档路径: {file_path}
分析需求: {analysis_requirements}

分析步骤:
1. 首先使用process_word_document工具处理文档
2. 根据需求选择合适的分析类型(summary/keywords/structure)
3. 使用analyze_document_content进行详细分析
4. 提供专业的分析结果和建议

请开始分析..."""
    
    return mcp


# 为了向后兼容，保留原有的函数名
process_word_document = process_word_document_func
analyze_document_content = analyze_document_content_func