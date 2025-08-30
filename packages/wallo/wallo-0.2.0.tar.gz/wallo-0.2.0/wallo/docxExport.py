"""DOCX export functionality for the Wallo application."""
import re
from typing import cast
from docx import Document  # pylint: disable=import-outside-toplevel
from docx.shared import RGBColor  # pylint: disable=import-outside-toplevel
from PySide6.QtCore import QObject  # pylint: disable=no-name-in-module
from PySide6.QtGui import QColor, QFont  # pylint: disable=no-name-in-module
from PySide6.QtWidgets import QMessageBox, QTextEdit # pylint: disable=no-name-in-module

class DocxExporter(QObject):
    """Handles DOCX export functionality with formatting preservation."""

    def exportToDocx(self, editor: QTextEdit, filename: str) -> None:
        """Export QTextEdit content to DOCX file with formatting preservation.

        Args:
            editor: The QTextEdit widget containing the content
            filename: The output filename for the DOCX file
        """
        try:
            self.exportWithPythonDocx(editor, filename)
        except Exception as error:
            QMessageBox.critical(cast(QTextEdit, self.parent()), 'Save Error', f"Failed to save document: {error}")


    def exportWithPythonDocx(self, editor: QTextEdit, filename: str) -> None:
        """Export content using python-docx library for better formatting preservation.

        Args:
            editor: The QTextEdit widget containing the content
            filename: The output filename for the DOCX file
        """
        document = Document()
        textDocument = editor.document()
        for blockNumber in range(textDocument.blockCount()):
            block = textDocument.findBlockByNumber(blockNumber)
            if not block.isValid():
                continue

            # Create a new paragraph for each block
            paragraph = document.add_paragraph()

            # Get the block iterator
            blockIterator = block.begin()

            # Process fragments within the block
            while not blockIterator.atEnd():
                fragment = blockIterator.fragment()
                if fragment.isValid():
                    fragmentText = fragment.text()
                    charFormat = fragment.charFormat()

                    # Create a run for this fragment
                    run = paragraph.add_run(fragmentText)

                    # Apply formatting based on character format
                    if charFormat.fontWeight() == QFont.Weight.Bold:
                        run.bold = True
                    if charFormat.fontItalic():
                        run.italic = True
                    if charFormat.fontUnderline():
                        run.underline = True

                    # Apply color if it's not default
                    color = charFormat.foreground().color()
                    if color.isValid() and color != QColor(0, 0, 0):  # Not default black
                        run.font.color.rgb = RGBColor(color.red(), color.green(), color.blue())

                blockIterator += 1

        document.save(filename)

    def cleanHtmlForDocx(self, html: str) -> str:
        """Clean HTML content for better DOCX conversion.

        Args:
            html: Raw HTML from QTextEdit

        Returns:
            Cleaned HTML suitable for DOCX conversion
        """
        # Convert Qt-specific color formatting to standard HTML
        html = re.sub(r'style="[^"]*color:\s*rgb\(([^)]+)\)[^"]*"',
                     lambda m: f'style="color: rgb({m.group(1)})"', html)

        # Convert Qt font-weight to standard HTML
        html = re.sub(r'style="[^"]*font-weight:\s*([^;]+)[^"]*"',
                     lambda m: f'style="font-weight: {m.group(1)}"', html)

        # Convert Qt font-style to standard HTML
        html = re.sub(r'style="[^"]*font-style:\s*([^;]+)[^"]*"',
                     lambda m: f'style="font-style: {m.group(1)}"', html)

        # Convert Qt text-decoration to standard HTML
        html = re.sub(r'style="[^"]*text-decoration:\s*([^;]+)[^"]*"',
                     lambda m: f'style="text-decoration: {m.group(1)}"', html)

        # Remove Qt-specific margin and padding that might interfere
        html = re.sub(r'margin-top:\s*[^;]+;?', '', html)
        html = re.sub(r'margin-bottom:\s*[^;]+;?', '', html)
        html = re.sub(r'margin-left:\s*[^;]+;?', '', html)
        html = re.sub(r'margin-right:\s*[^;]+;?', '', html)

        # Clean up empty style attributes
        html = re.sub(r'style=""', '', html)
        html = re.sub(r'style="[;\s]*"', '', html)

        # Ensure proper HTML structure
        if not html.startswith('<html'):
            html = f'<html><head><meta charset="UTF-8"></head><body>{html}</body></html>'
        return html
